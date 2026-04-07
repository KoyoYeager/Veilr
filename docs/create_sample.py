"""Veilr デモ用 Word ファイル — 1枚のスクリーンショットで魅力が伝わる高密度デザイン"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

# カラーパレット
RED = RGBColor(0xFF, 0x00, 0x00)
DARK_RED = RGBColor(0xC0, 0x30, 0x30)
BRIGHT_RED = RGBColor(0xFF, 0x33, 0x33)
MATERIAL_RED = RGBColor(0xD3, 0x2F, 0x2F)
CRIMSON = RGBColor(0xDC, 0x14, 0x3C)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
GRAY = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def shade(cell, hex_color: str):
    tcPr = cell._element.get_or_add_tcPr()
    tcPr.append(
        tcPr.makeelement(
            qn("w:shd"), {qn("w:fill"): hex_color, qn("w:val"): "clear"}
        )
    )


def run(para, text, color=BLACK, bold=False, size=11, font="游ゴシック"):
    r = para.add_run(text)
    r.font.color.rgb = color
    r.font.bold = bold
    r.font.size = Pt(size)
    r.font.name = font
    r._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    return r


def set_row_height(row, cm):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = trPr.makeelement(
        qn("w:trHeight"), {qn("w:val"): str(int(cm * 567)), qn("w:hRule"): "exact"}
    )
    trPr.append(trHeight)


def set_spacing(para, before=0, after=0):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after = Pt(after)


def main():
    doc = Document()

    # ページ設定 — A4, 狭い余白で密度を上げる
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "游ゴシック"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "游ゴシック")

    # ━━━━━━ タイトル ━━━━━━
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, 4, 2)
    run(title, "期末試験対策  情報セキュリティ完全まとめ", BLACK, bold=True, size=22)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, 0, 8)
    run(subtitle, "赤い文字 ＝ 答え。赤シートで隠して暗記しよう", GRAY, size=10)

    doc.add_paragraph()  # 小スペース

    # ━━━━━━ セクションA: 超重要キーワード（大きな赤文字） ━━━━━━
    h = doc.add_paragraph()
    set_spacing(h, 0, 4)
    run(h, "■ 最重要キーワード", DARK_GRAY, bold=True, size=13)

    # 3カラムテーブルで大きな赤キーワードを並べる
    keywords = [
        ("機密性", "Confidentiality"),
        ("完全性", "Integrity"),
        ("可用性", "Availability"),
    ]
    kt = doc.add_table(rows=2, cols=3)
    kt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, (ja, en) in enumerate(keywords):
        # 上段: 日本語（大きな赤文字）
        cell = kt.rows[0].cells[j]
        shade(cell, "FFF5F5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, 6, 0)
        run(p, ja, RED, bold=True, size=20)
        # 下段: 英語
        cell2 = kt.rows[1].cells[j]
        shade(cell2, "FFF5F5")
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p2, 0, 6)
        run(p2, en, CRIMSON, size=9)

    # 説明文
    desc = doc.add_paragraph()
    set_spacing(desc, 4, 2)
    run(desc, "情報セキュリティの3要素（CIAトライアド）。この3つを守ることが全ての基本。")

    # ━━━━━━ セクションB: 穴埋め問題 ━━━━━━
    doc.add_paragraph()
    h2 = doc.add_paragraph()
    set_spacing(h2, 0, 4)
    run(h2, "■ 穴埋めテスト", DARK_GRAY, bold=True, size=13)

    qa = [
        ("不正アクセスを検知するシステムは ", "IDS", "（Intrusion Detection System）"),
        ("通信を暗号化するプロトコルは ", "TLS/SSL", ""),
        ("パスワード保存時に付加するランダム値は ", "ソルト（salt）", ""),
        ("SQL注入を防ぐには ", "プリペアドステートメント", " を使う"),
        ("公開鍵暗号の代表的アルゴリズムは ", "RSA", " である"),
        ("脆弱性トップ10プロジェクトは ", "OWASP Top 10", ""),
        ("ゼロトラストの原則は ", "何も信頼せず全てを検証", " すること"),
        ("多要素認証(MFA)は ", "知識・所持・生体", " の組合せ"),
    ]

    for i, (q, a, tail) in enumerate(qa, 1):
        p = doc.add_paragraph()
        set_spacing(p, 1, 1)
        run(p, f"({i:02d}) {q}", BLACK, size=10)
        run(p, a, RED, bold=True, size=11)
        if tail:
            run(p, tail, BLACK, size=10)

    # ━━━━━━ セクションC: 攻撃と対策テーブル ━━━━━━
    doc.add_paragraph()
    h3 = doc.add_paragraph()
    set_spacing(h3, 0, 4)
    run(h3, "■ 攻撃手法 × 対策 一覧", DARK_GRAY, bold=True, size=13)

    # サブ注釈
    note = doc.add_paragraph()
    set_spacing(note, 0, 4)
    run(note, "※ 攻撃名・対策とも赤字。Veilrの色ファミリー検出で濃淡の違う赤もまとめて消えます", GRAY, size=8)

    rows_data = [
        ("フィッシング", "偽サイトで認証情報を窃取", "URL確認 + MFA導入", RED),
        ("SQLインジェクション", "入力値にSQL文を注入", "プリペアドステートメント", DARK_RED),
        ("XSS", "悪意あるスクリプトを埋込", "出力エスケープ + CSP", BRIGHT_RED),
        ("CSRF", "意図しないリクエスト送信", "CSRFトークン検証", MATERIAL_RED),
        ("ブルートフォース", "総当たりでパスワード試行", "アカウントロック + MFA", CRIMSON),
        ("中間者攻撃(MITM)", "通信を傍受・改ざん", "TLS/HTTPS必須化", RGBColor(0xE8, 0x00, 0x00)),
    ]

    t = doc.add_table(rows=len(rows_data) + 1, cols=3)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー
    for j, text in enumerate(["攻撃手法", "概要", "対策"]):
        c = t.rows[0].cells[j]
        shade(c, "1A1A2E")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, 3, 3)
        run(p, text, WHITE, bold=True, size=10)

    # データ
    for i, (attack, desc_text, defense, color) in enumerate(rows_data, 1):
        # 攻撃名 — 各行で異なる赤色（色ファミリーのデモ）
        c0 = t.rows[i].cells[0]
        p0 = c0.paragraphs[0]
        set_spacing(p0, 2, 2)
        run(p0, attack, color, bold=True, size=10)

        # 概要 — 黒
        c1 = t.rows[i].cells[1]
        p1 = c1.paragraphs[0]
        set_spacing(p1, 2, 2)
        run(p1, desc_text, BLACK, size=9)

        # 対策 — 赤
        c2 = t.rows[i].cells[2]
        p2 = c2.paragraphs[0]
        set_spacing(p2, 2, 2)
        run(p2, defense, RED, bold=True, size=9)

    # ━━━━━━ セクションD: 暗号方式比較 ━━━━━━
    doc.add_paragraph()
    h4 = doc.add_paragraph()
    set_spacing(h4, 0, 4)
    run(h4, "■ 暗号方式の比較", DARK_GRAY, bold=True, size=13)

    crypto = [
        ["方式", "鍵", "代表例", "速度", "用途"],
        ["共通鍵", "1つ", "AES-256", "高速", "データ暗号化"],
        ["公開鍵", "2つ", "RSA-2048", "低速", "鍵交換・署名"],
        ["ハイブリッド", "両方", "TLS 1.3", "実用的", "Web通信"],
    ]

    t2 = doc.add_table(rows=4, cols=5)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER

    for j, text in enumerate(crypto[0]):
        c = t2.rows[0].cells[j]
        shade(c, "1A1A2E")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, 3, 3)
        run(p, text, WHITE, bold=True, size=9)

    red_cols = {2, 3}  # 代表例と速度を赤
    for i in range(1, 4):
        for j in range(5):
            c = t2.rows[i].cells[j]
            p = c.paragraphs[0]
            set_spacing(p, 2, 2)
            if j in red_cols:
                shade(c, "FFF8F8")
                run(p, crypto[i][j], RED, bold=True, size=9)
            else:
                run(p, crypto[i][j], BLACK, size=9)

    # ━━━━━━ セクションE: 重要公式 ━━━━━━
    doc.add_paragraph()
    h5 = doc.add_paragraph()
    set_spacing(h5, 0, 4)
    run(h5, "■ 覚えるべき数値", DARK_GRAY, bold=True, size=13)

    formulas = [
        ("AES鍵長", "128 / 192 / 256 ビット"),
        ("RSA推奨鍵長", "2048 ビット以上"),
        ("TLS最新バージョン", "1.3（2018年策定）"),
        ("パスワード最小長（NIST推奨）", "8文字以上、できれば15文字以上"),
        ("HTTPSデフォルトポート", "443"),
        ("SSH デフォルトポート", "22"),
    ]

    ft = doc.add_table(rows=len(formulas), cols=2)
    ft.style = "Table Grid"
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, (label, value) in enumerate(formulas):
        c0 = ft.rows[i].cells[0]
        shade(c0, "F5F5F5")
        p0 = c0.paragraphs[0]
        set_spacing(p0, 2, 2)
        run(p0, label, DARK_GRAY, bold=True, size=9)

        c1 = ft.rows[i].cells[1]
        p1 = c1.paragraphs[0]
        set_spacing(p1, 2, 2)
        run(p1, value, RED, bold=True, size=10)

    # ━━━━━━ フッター ━━━━━━
    doc.add_paragraph()
    ft_banner = doc.add_table(rows=1, cols=1)
    ft_banner.alignment = WD_TABLE_ALIGNMENT.CENTER
    fc = ft_banner.rows[0].cells[0]
    shade(fc, "F0F0F0")
    fp = fc.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(fp, 6, 6)
    run(fp, "Veilr", DARK_GRAY, bold=True, size=12)
    run(fp, "  — 画面上の赤い文字を隠す・消す。デジタル赤シート for Windows", GRAY, size=9)

    output = "C:/workspace/Veilr/docs/Veilr-sample.docx"
    doc.save(output)
    print(f"Saved: {output}")


if __name__ == "__main__":
    main()
