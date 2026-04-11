"""Veilr記事用デモWordファイル — 1ページに凝縮"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles["Normal"]
style.font.name = "游ゴシック"
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.3

for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xE0, 0x20, 0x20)


def add_mixed(parts: list[tuple[str, RGBColor, bool]], size: int = 12, indent: bool = False) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    for text, color, bold in parts:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold


# === タイトル ===
h = doc.add_heading("", level=1)
run = h.add_run("Python & Web 基礎チェック")
run.font.size = Pt(20)
run.font.color.rgb = BLACK
run.font.bold = True
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run("赤い文字 ＝ 答え。赤シートで隠して暗記しよう")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(8)

# === セクション1: 穴埋め ===
h = doc.add_heading("", level=2)
run = h.add_run("■ 穴埋めテスト")
run.font.size = Pt(14)
run.font.color.rgb = BLACK
run.font.bold = True
h.paragraph_format.space_before = Pt(4)
h.paragraph_format.space_after = Pt(2)

blanks = [
    [("(1) Pythonの変数は ", BLACK, False), ("型宣言が不要", RED, True), (" で、実行時に型が決まる", BLACK, False)],
    [("(2) 例外処理は ", BLACK, False), ("try-except", RED, True), (" 構文で書き、後処理は ", BLACK, False), ("finally", RED, True), (" に置く", BLACK, False)],
    [("(3) リスト内包表記の構文は ", BLACK, False), ("[式 for 変数 in イテラブル]", RED, True)],
    [("(4) HTTPステータス 404 は ", BLACK, False), ("Not Found", RED, True), ("、500 は ", BLACK, False), ("Internal Server Error", RED, True)],
    [("(5) RESTful APIで、リソース作成に使うHTTPメソッドは ", BLACK, False), ("POST", RED, True)],
]

for parts in blanks:
    add_mixed(parts, size=12, indent=True)

# === セクション2: フレームワーク比較表 ===
h = doc.add_heading("", level=2)
run = h.add_run("■ Webフレームワーク比較")
run.font.size = Pt(14)
run.font.color.rgb = BLACK
run.font.bold = True
h.paragraph_format.space_before = Pt(8)
h.paragraph_format.space_after = Pt(2)

p = doc.add_paragraph()
run = p.add_run("※ 赤文字が答え。消去モードで消すと空欄の暗記表になる")
run.font.size = Pt(10)
run.font.color.rgb = GRAY
p.paragraph_format.space_after = Pt(4)

# ヘッダー + データ: (言語, FW名, 特徴, 用途) — FW名と特徴が赤
rows_data = [
    ("言語", "フレームワーク", "特徴", "用途"),
    ("Python", "Django", "フルスタック・管理画面付き", "業務系Webアプリ"),
    ("Python", "FastAPI", "非同期・自動API docs", "REST API"),
    ("JavaScript", "Next.js", "SSR / SSG対応", "フルスタックReact"),
    ("Ruby", "Rails", "CoC・DRY原則", "高速プロトタイピング"),
    ("Go", "Gin", "高速・軽量", "マイクロサービス"),
]

table = doc.add_table(rows=len(rows_data), cols=4)
table.style = "Table Grid"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 答えセル: FW名(col1)と特徴(col2) が赤
answer_cols = {1, 2}

for i, (c1, c2, c3, c4) in enumerate(rows_data):
    for j, text in enumerate([c1, c2, c3, c4]):
        cell = table.rows[i].cells[j]
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(11)
        run.font.name = "游ゴシック"
        if i == 0:
            run.font.bold = True
            run.font.color.rgb = BLACK
        elif j in answer_cols:
            run.font.color.rgb = RED
            run.font.bold = True
        else:
            run.font.color.rgb = BLACK

for row in table.rows:
    row.cells[0].width = Cm(2.8)
    row.cells[1].width = Cm(3.5)
    row.cells[2].width = Cm(5)
    row.cells[3].width = Cm(4.5)

# === セクション3: 一問一答 ===
h = doc.add_heading("", level=2)
run = h.add_run("■ 一問一答")
run.font.size = Pt(14)
run.font.color.rgb = BLACK
run.font.bold = True
h.paragraph_format.space_before = Pt(8)
h.paragraph_format.space_after = Pt(2)

qa = [
    ("Q1. Gitでブランチを作成して切り替えるコマンドは？", "git checkout -b <name>"),
    ("Q2. CSSで要素を横並びにするプロパティは？", "display: flex"),
    ("Q3. SQLで重複を除いて取得するキーワードは？", "DISTINCT"),
    ("Q4. Dockerでイメージからコンテナを起動するコマンドは？", "docker run"),
]

for q, a in qa:
    add_mixed([
        (f"  {q}", BLACK, False),
    ], size=12)
    add_mixed([
        ("　　→ ", BLACK, False),
        (a, RED, True),
    ], size=12)

# === フッター ===
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Veilr — 画面の色を隠せるツール for Windows")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

output = "C:/workspace/Veilr/docs/article-demo.docx"
doc.save(output)
print(f"saved: {output}")
